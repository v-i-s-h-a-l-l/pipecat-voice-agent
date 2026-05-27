from pathlib import Path
from typing import Optional

import docx

from llama_index.core import Document
from llama_index.core.readers.base import BaseReader

FILENAME_TO_DOCTYPE = {
    "Complete_Menu": "menu",
    "customerexperience": "qa",
    "Restaurant_Grand_Chennai": "fact",
    "TABLE_BOOKING_SYSTEM": "policy",
}


class StructuralDocxReader(BaseReader):

    def load_data(
        self,
        file: Path,
        extra_info: Optional[dict] = None,
    ) -> list[Document]:

        stem = file.stem.replace(" ", "_").replace("-", "_")

        doc_type = FILENAME_TO_DOCTYPE.get(stem, "fact")

        word_doc = docx.Document(file)

        documents = []

        current_section = "root"

        for para in word_doc.paragraphs:

            text = para.text.strip()

            if not text:
                continue

            style = para.style.name if para.style else ""

            heading_level = 0

            if "Heading 1" in style:
                heading_level = 1
                current_section = text

            elif "Heading 2" in style:
                heading_level = 2
                current_section = text

            documents.append(
                Document(
                    text=text,
                    metadata={
                        "doc_type": doc_type,
                        "source_file": file.name,
                        "section": current_section,
                        "heading_level": heading_level,
                    },
                )
            )

        return documents